from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
import tempfile
import os
from pathlib import Path

from models import ExportRequest, ExportFormat

router = APIRouter(prefix="/api/export", tags=["export"])


def markdown_to_pdf(markdown_content: str, output_path: str):
    """Convert markdown to PDF using md2pdf."""
    try:
        from md2pdf.core import md2pdf
        md2pdf(output_path, md_content=markdown_content)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate PDF: {str(e)}"
        )


def markdown_to_docx(markdown_content: str, output_path: str):
    """Convert markdown to DOCX using python-docx with basic markdown parsing."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import re

        doc = Document()

        # Split content by lines
        lines = markdown_content.split('\n')

        for line in lines:
            line = line.rstrip()

            # Skip empty lines
            if not line:
                doc.add_paragraph()
                continue

            # Headers
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)

            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')

            # Regular paragraph
            else:
                p = doc.add_paragraph(line)

        doc.save(output_path)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate DOCX: {str(e)}"
        )


@router.post("")
async def export_markdown(request: ExportRequest):
    """
    Export markdown content to PDF or DOCX format.

    Args:
        request: ExportRequest containing format and markdown content

    Returns:
        FileResponse: The generated file

    Raises:
        HTTPException: 500 if export fails
    """
    # Create temporary file
    suffix = f".{request.format.value}"
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    temp_path = temp_file.name
    temp_file.close()

    try:
        # Generate file based on format
        if request.format == ExportFormat.PDF:
            markdown_to_pdf(request.markdown, temp_path)
            media_type = "application/pdf"
            filename = "export.pdf"
        elif request.format == ExportFormat.DOCX:
            markdown_to_docx(request.markdown, temp_path)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = "export.docx"
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")

        # Return file response with cleanup
        return FileResponse(
            path=temp_path,
            media_type=media_type,
            filename=filename,
            background=lambda: os.unlink(temp_path)  # Cleanup temp file after response
        )

    except HTTPException:
        # Cleanup on error
        if os.path.exists(temp_path):
            os.unlink(temp_path)
        raise
    except Exception as e:
        # Cleanup on unexpected error
        if os.path.exists(temp_path):
            os.unlink(temp_path)
        raise HTTPException(
            status_code=500,
            detail=f"Export failed: {str(e)}"
        )
